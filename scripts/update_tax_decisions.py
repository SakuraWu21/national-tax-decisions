#!/usr/bin/env python3
"""全国税务处理及行政处罚决定书：发现、核验、去重、持久化。

设计原则：
1. 搜索结果只能作为候选，写入前必须实际访问正文或附件。
2. 不猜测缺失字段；无法确认时留空。
3. Excel 与 CSV 都是历史数据输入，合并后再执行增量更新。
4. 冲突不静默覆盖，统一写入“备注”。
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import io
import json
import os
import re
import shutil
import subprocess
import sys
import time
import unicodedata
import zipfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Iterable
from urllib.parse import parse_qs, quote_plus, unquote, urljoin, urlparse, urlunparse
from xml.etree import ElementTree
from zoneinfo import ZoneInfo

import requests
from bs4 import BeautifulSoup
from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.formatting.rule import FormulaRule
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation
from pypdf import PdfReader
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data"
XLSX_PATH = DATA_DIR / "全国税务处理及行政处罚决定书汇总.xlsx"
CSV_PATH = DATA_DIR / "全国税务处理及行政处罚决定书汇总.csv"
AUDIT_PATH = DATA_DIR / "candidate_audit.jsonl"
RETRY_QUEUE_PATH = DATA_DIR / "retry_queue.json"
SEED_PATH = ROOT / "config" / "seed_urls.txt"
OFFICIAL_INDEX_PATH = ROOT / "config" / "official_indexes.txt"
PUBLIC_DATA_DIR = ROOT / "public" / "data"
PUBLIC_DOWNLOAD_DIR = ROOT / "public" / "downloads"
PUBLIC_JSON_PATH = PUBLIC_DATA_DIR / "tax-decisions.json"
PUBLIC_STATUS_PATH = PUBLIC_DATA_DIR / "update-status.json"
PUBLIC_SOURCE_HEALTH_PATH = PUBLIC_DATA_DIR / "source-health.json"
PUBLIC_XLSX_PATH = PUBLIC_DOWNLOAD_DIR / "全国税务处理及行政处罚决定书汇总.xlsx"
PUBLIC_CSV_PATH = PUBLIC_DOWNLOAD_DIR / "全国税务处理及行政处罚决定书汇总.csv"

TZ_NAME = "Asia/Shanghai"
PUBLIC_SCHEDULE = "每日 12:07 主任务；12:37 补偿任务（北京时间）"
RECENT_LOOKBACK_DAYS = 7
OFFICIAL_INDEX_LOOKBACK_DAYS = 14


def now_in_project_timezone() -> datetime:
    """所有环境统一使用项目约定的北京时间，避免受执行主机时区影响。"""
    return datetime.now(ZoneInfo(TZ_NAME))


RUN_NOW = now_in_project_timezone()
TODAY = RUN_NOW.date()

FIELDS = [
    "序号", "案件组ID", "文书唯一ID", "省份", "城市", "区县", "发布机关", "稽查机构",
    "当事人名称", "统一社会信用代码", "法定代表人", "文书类型", "决定书文号",
    "关联处理决定书文号", "关联处罚决定书文号", "主要违法事实", "涉及税种",
    "追缴税款金额", "滞纳金金额", "罚款金额", "没收违法所得金额", "处理或处罚结果",
    "决定书作出日期", "官方发布日期", "第三方收录日期", "首次发现日期", "最后核验日期",
    "公开完整度", "来源级别", "核验状态", "官方原文链接", "附件链接", "备用来源链接",
    "页面标题", "页面当前状态", "备注",
]

LOG_FIELDS = [
    "运行日期和时间", "检索时间范围", "检索页面数量", "新增完整文书数量",
    "新增文号线索数量", "更新旧记录数量", "重复记录数量", "失效链接数量",
    "待人工核验数量", "本次新增当事人名单", "运行是否成功", "错误摘要", "运行说明",
]

DOC_TYPES = ("税务处理决定书", "税务行政处罚决定书")
EXCLUDED_TITLES = (
    "税务行政处罚事项告知书", "税务处理事项告知书", "税务检查通知书", "税务稽查通知书",
    "责令限期改正通知书", "催告书", "听证通知书", "听证权利告知书", "欠税公告",
)
# 这类页面可能在正文中引用既往决定书文号，但页面本身公开的是另一份文书。
# 只有标题同时明确包含目标决定书类型时，才允许继续核验正文或附件。
EXCLUDED_ONLY_TITLE_MARKERS = ("税务事项通知书",)
OFFICIAL_SUFFIXES = (".chinatax.gov.cn", ".gov.cn")
NORMAL_STATES = {"正常", "附件正常"}
SOURCE_PRIORITY = {"第三方待核验": 1, "政府公开平台": 2, "税务机关官网": 3}
HTTPS_UPGRADE_HOSTS = {"neimenggu.chinatax.gov.cn"}
INVALID_LEGAL_REP_TOKENS = (
    "已", "法院", "涤除", "无法", "联系", "失联",
    "委托代理人", "单位公章", "身份证", "材料办理", "或者股东", "个人账户",
)
INVALID_PARTY_MARKERS = ("该地址也是", "同时,", "同时，", "同时该地址")

MONEY_FIELDS = {"追缴税款金额", "滞纳金金额", "罚款金额", "没收违法所得金额"}
DATE_FIELDS = {"决定书作出日期", "官方发布日期", "第三方收录日期", "首次发现日期", "最后核验日期"}
LINK_FIELDS = {"官方原文链接", "附件链接", "备用来源链接"}
LONG_TEXT_FIELDS = {"主要违法事实", "处理或处罚结果", "备注", "页面标题"}

JSON_FIELD_MAP = {
    "id": "文书唯一ID",
    "caseGroupId": "案件组ID",
    "province": "省份",
    "city": "城市",
    "district": "区县",
    "issuingAuthority": "发布机关",
    "inspectionAuthority": "稽查机构",
    "partyName": "当事人名称",
    "unifiedSocialCreditCode": "统一社会信用代码",
    "legalRepresentative": "法定代表人",
    "documentType": "文书类型",
    "documentNumber": "决定书文号",
    "relatedTreatmentDocumentNumber": "关联处理决定书文号",
    "relatedPenaltyDocumentNumber": "关联处罚决定书文号",
    "violationFacts": "主要违法事实",
    "taxTypes": "涉及税种",
    "recoveredTaxAmount": "追缴税款金额",
    "lateFeeAmount": "滞纳金金额",
    "fineAmount": "罚款金额",
    "confiscatedIncomeAmount": "没收违法所得金额",
    "result": "处理或处罚结果",
    "decisionDate": "决定书作出日期",
    "officialPublishDate": "官方发布日期",
    "thirdPartyPublishDate": "第三方收录日期",
    "firstDiscoveredDate": "首次发现日期",
    "lastVerifiedDate": "最后核验日期",
    "completeness": "公开完整度",
    "sourceLevel": "来源级别",
    "verificationStatus": "核验状态",
    "officialUrl": "官方原文链接",
    "attachmentUrl": "附件链接",
    "backupUrl": "备用来源链接",
    "pageTitle": "页面标题",
    "pageStatus": "页面当前状态",
    "notes": "备注",
}

PROVINCES = [
    "北京市", "天津市", "上海市", "重庆市", "河北省", "山西省", "辽宁省", "吉林省",
    "黑龙江省", "江苏省", "浙江省", "安徽省", "福建省", "江西省", "山东省", "河南省",
    "湖北省", "湖南省", "广东省", "海南省", "四川省", "贵州省", "云南省", "陕西省",
    "甘肃省", "青海省", "台湾省", "内蒙古自治区", "广西壮族自治区", "西藏自治区",
    "宁夏回族自治区", "新疆维吾尔自治区", "香港特别行政区", "澳门特别行政区",
]
PROVINCE_ALIASES = {
    "北京": "北京市", "天津": "天津市", "上海": "上海市", "重庆": "重庆市",
    "河北": "河北省", "山西": "山西省", "辽宁": "辽宁省", "吉林": "吉林省",
    "黑龙江": "黑龙江省", "江苏": "江苏省", "浙江": "浙江省", "安徽": "安徽省",
    "福建": "福建省", "江西": "江西省", "山东": "山东省", "河南": "河南省",
    "湖北": "湖北省", "湖南": "湖南省", "广东": "广东省", "海南": "海南省",
    "四川": "四川省", "贵州": "贵州省", "云南": "云南省", "陕西": "陕西省",
    "甘肃": "甘肃省", "青海": "青海省", "内蒙古": "内蒙古自治区",
    "广西": "广西壮族自治区", "西藏": "西藏自治区", "宁夏": "宁夏回族自治区",
    "新疆": "新疆维吾尔自治区",
}
DOMAIN_PROVINCES = {
    "beijing": "北京市", "tianjin": "天津市", "hebei": "河北省", "shanxi": "山西省",
    "neimenggu": "内蒙古自治区", "liaoning": "辽宁省", "dalian": "辽宁省", "jilin": "吉林省",
    "heilongjiang": "黑龙江省", "shanghai": "上海市", "jiangsu": "江苏省", "zhejiang": "浙江省",
    "ningbo": "浙江省", "anhui": "安徽省", "fujian": "福建省", "xiamen": "福建省",
    "jiangxi": "江西省", "shandong": "山东省", "qingdao": "山东省", "henan": "河南省",
    "hubei": "湖北省", "hunan": "湖南省", "guangdong": "广东省", "shenzhen": "广东省",
    "guangxi": "广西壮族自治区", "hainan": "海南省", "chongqing": "重庆市", "sichuan": "四川省",
    "guizhou": "贵州省", "yunnan": "云南省", "xizang": "西藏自治区", "shaanxi": "陕西省",
    "gansu": "甘肃省", "qinghai": "青海省", "ningxia": "宁夏回族自治区",
    "xinjiang": "新疆维吾尔自治区",
}
CORE_SOURCE_PROVINCES = ("福建省", "河南省", "贵州省", "广西壮族自治区", "广东省", "江苏省", "浙江省")

TAX_NAMES = [
    "增值税", "企业所得税", "个人所得税", "消费税", "城市维护建设税", "城建税",
    "房产税", "城镇土地使用税", "土地增值税", "印花税", "资源税", "契税",
    "车船税", "教育费附加", "地方教育附加", "文化事业建设费",
]
ENTITY_SUFFIXES = (
    "有限责任公司", "股份有限公司", "有限公司", "公司", "合作社", "事务所",
    "经营部", "商行", "商贸部", "厂", "店", "中心",
)

SEARCH_QUERIES_FULL = [
    'site:chinatax.gov.cn "税务处理决定书" "2026"',
    'site:chinatax.gov.cn "税务行政处罚决定书" "2026"',
    'site:chinatax.gov.cn "税务处理决定书" "公告送达" 2026',
    'site:chinatax.gov.cn "税务行政处罚决定书" "公告送达" 2026',
    'site:chinatax.gov.cn "税稽处〔2026〕"',
    'site:chinatax.gov.cn "税稽罚〔2026〕"',
    'site:chinatax.gov.cn "税一稽处〔2026〕"',
    'site:chinatax.gov.cn "税一稽罚〔2026〕"',
    'site:chinatax.gov.cn "税二稽处〔2026〕"',
    'site:chinatax.gov.cn "税二稽罚〔2026〕"',
    'site:gov.cn "税务处理决定书" "2026"',
    'site:gov.cn "税务行政处罚决定书" "2026"',
    '"稽查局" "税务处理决定书" "2026" "公告"',
    '"稽查局" "税务行政处罚决定书" "2026" "公告"',
    '"决定书全文" "税务处理决定书" 2026',
    '"国家税务总局" "公告送达" "处理决定书" 2026',
]

SEARCH_QUERIES_RECENT = [
    'site:chinatax.gov.cn "税务处理决定书" "公告送达"',
    'site:chinatax.gov.cn "税务行政处罚决定书" "公告送达"',
    'site:chinatax.gov.cn "税稽处"',
    'site:chinatax.gov.cn "税稽罚"',
    'site:gov.cn "税务处理决定书"',
    'site:gov.cn "税务行政处罚决定书"',
    '"稽查局" "税务处理决定书" "决定书全文"',
    '"稽查局" "税务行政处罚决定书" "决定书全文"',
]

DOC_NO_PATTERN = re.compile(
    r"(?P<no>[\u4e00-\u9fffA-Za-z]{1,12}税[\u4e00-\u9fffA-Za-z]{0,8}"
    r"(?:处|罚)(?:字)?\s*[〔﹝［\[\(（]?\s*20\d{2}\s*[〕﹞］\]\)）]?\s*"
    r"(?:第\s*)?\d{1,6}\s*号)",
    re.I,
)
DATE_PATTERN = re.compile(r"(20\d{2})\s*[年./\-]\s*(\d{1,2})\s*[月./\-]\s*(\d{1,2})\s*日?")
USCC_PATTERN = re.compile(r"(?<![A-Z0-9])([0-9A-HJ-NPQRTUWXY]{18})(?![A-Z0-9])", re.I)
AMOUNT_PATTERN = r"([0-9][0-9,，]*(?:\.[0-9]{1,2})?)\s*元"


@dataclass
class SearchHit:
    url: str
    title: str = ""
    provider: str = ""
    query: str = ""


@dataclass
class FetchResult:
    url: str
    final_url: str
    ok: bool
    status_code: int | None
    content_type: str
    content: bytes
    text: str
    title: str
    state: str
    error: str = ""
    attachments: list[dict] = field(default_factory=list)


def console(message: str) -> None:
    print(message, flush=True)


def new_session() -> requests.Session:
    session = requests.Session()
    retry = Retry(
        total=3,
        connect=3,
        read=3,
        status=3,
        backoff_factor=1.0,
        status_forcelist=(429, 500, 502, 503, 504),
        allowed_methods=("GET", "HEAD"),
        respect_retry_after_header=True,
    )
    session.mount("https://", HTTPAdapter(max_retries=retry, pool_connections=20, pool_maxsize=20))
    session.mount("http://", HTTPAdapter(max_retries=retry, pool_connections=20, pool_maxsize=20))
    session.headers.update({
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
            "(KHTML, like Gecko) Chrome/124.0 Safari/537.36"
        ),
        "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.5",
        "Accept": "text/html,application/xhtml+xml,application/pdf,application/xml;q=0.9,*/*;q=0.8",
    })
    return session


def clean_text(value: object) -> str:
    if value is None:
        return ""
    text = unicodedata.normalize("NFKC", str(value))
    text = text.replace("\u00a0", " ").replace("\u3000", " ")
    return re.sub(r"[ \t]+", " ", text).strip()


def compact_multiline(value: object) -> str:
    text = clean_text(value)
    text = re.sub(r"\s*\n\s*", "\n", text)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def excel_safe(value: object) -> object:
    if isinstance(value, str) and value[:1] in ("=", "+", "-", "@"):
        return "'" + value
    return value


def normalize_url(url: str) -> str:
    url = clean_text(url)
    if not url:
        return ""
    try:
        parsed = urlparse(url)
        if parsed.netloc.endswith("bing.com") and parsed.path.startswith("/ck/a"):
            qs = parse_qs(parsed.query)
            encoded = (qs.get("u") or [""])[0]
            if encoded.startswith("a1"):
                import base64
                encoded += "=" * (-len(encoded) % 4)
                url = base64.urlsafe_b64decode(encoded[2:]).decode("utf-8", "ignore")
                parsed = urlparse(url)
        host = parsed.netloc.lower().split("@")[-1]
        if host.startswith("www."):
            host = host[4:]
        path = re.sub(r"/{2,}", "/", unquote(parsed.path or "/"))
        query_items = []
        for key, values in parse_qs(parsed.query, keep_blank_values=True).items():
            if key.lower() in {"utm_source", "utm_medium", "utm_campaign", "utm_term", "utm_content", "spm"}:
                continue
            for val in values:
                query_items.append((key, val))
        from urllib.parse import urlencode
        query = urlencode(sorted(query_items), doseq=True)
        scheme = parsed.scheme.lower() or "https"
        if scheme == "http" and host in HTTPS_UPGRADE_HOSTS:
            scheme = "https"
        return urlunparse((scheme, host, path.rstrip("/") or "/", "", query, ""))
    except Exception:
        return url


def is_official_url(url: str) -> bool:
    host = urlparse(url).netloc.lower().split(":")[0]
    return any(host == suffix[1:] or host.endswith(suffix) for suffix in OFFICIAL_SUFFIXES)


def source_level(url: str) -> str:
    host = urlparse(url).netloc.lower()
    if "chinatax.gov.cn" in host:
        return "税务机关官网"
    if host.endswith(".gov.cn") or host == "gov.cn":
        return "政府公开平台"
    return "第三方待核验"


def normalized_doc_no(value: str) -> str:
    text = clean_text(value)
    # 页面常把“决定书/公告送达”等文本与文号粘连；这些词不是文号的一部分。
    if "〔" in text or "[" in text or "（" in text or "(" in text:
        text = re.sub(r"^.*(?:决定书|送达|公告|下列单位|附件)", "", text)
    table = str.maketrans({"[": "〔", "(": "〔", "（": "〔", "﹝": "〔", "]": "〕", ")": "〕", "）": "〕", "﹞": "〕"})
    text = text.translate(table)
    text = re.sub(r"\s+", "", text)
    text = re.sub(r"〔(20\d{2})〕?", r"〔\1〕", text)
    return text


def doc_type_from_number(number: str, context: str = "") -> str:
    no = normalized_doc_no(number)
    prefix = no.split("〔", 1)[0]
    if re.search(r"罚(?:字)?$", prefix):
        return "税务行政处罚决定书"
    if re.search(r"处(?:字)?$", prefix):
        return "税务处理决定书"
    if "行政处罚决定书" in context:
        return "税务行政处罚决定书"
    return "税务处理决定书"


def parse_money(value: object) -> float | None:
    if value is None or value == "":
        return None
    if isinstance(value, (int, float)):
        return float(value)
    text = clean_text(value).replace(",", "").replace("，", "")
    try:
        return float(text)
    except ValueError:
        return None


def money_string(value: object) -> str:
    number = parse_money(value)
    if number is None:
        return ""
    return f"{number:.2f}".rstrip("0").rstrip(".")


def iso_date(value: object) -> str:
    if value is None or value == "":
        return ""
    if isinstance(value, datetime):
        return value.date().isoformat()
    if isinstance(value, date):
        return value.isoformat()
    text = clean_text(value)
    match = DATE_PATTERN.search(text)
    if match:
        try:
            return date(*map(int, match.groups())).isoformat()
        except ValueError:
            return ""
    match = re.search(r"(20\d{2})-(\d{1,2})-(\d{1,2})", text)
    if match:
        try:
            return date(*map(int, match.groups())).isoformat()
        except ValueError:
            return ""
    return ""


def hash_id(prefix: str, value: str, size: int = 16) -> str:
    digest = hashlib.sha256(value.encode("utf-8", "ignore")).hexdigest()[:size]
    return f"{prefix}-{digest}"


def build_unique_id(record: dict) -> str:
    doc_type = clean_text(record.get("文书类型"))
    doc_no = normalized_doc_no(record.get("决定书文号", ""))
    if doc_no:
        return hash_id("DOC", f"{doc_type}|{doc_no}")
    key = "|".join([
        clean_text(record.get("当事人名称")),
        doc_type,
        iso_date(record.get("决定书作出日期")),
        normalize_url(record.get("官方原文链接") or record.get("备用来源链接") or ""),
    ])
    return hash_id("DOC", key)


def decision_year(record: dict) -> str:
    doc_no = normalized_doc_no(record.get("决定书文号", ""))
    match = re.search(r"20\d{2}", doc_no)
    if match:
        return match.group(0)
    value = iso_date(record.get("决定书作出日期")) or iso_date(record.get("官方发布日期"))
    return value[:4] if value else ""


def build_group_id(record: dict) -> str:
    key = "|".join([
        clean_text(record.get("当事人名称")),
        clean_text(record.get("发布机关") or record.get("稽查机构")),
        decision_year(record),
    ])
    return hash_id("CASE", key)


def discover_bing_rss(session: requests.Session, query: str) -> list[SearchHit]:
    url = f"https://www.bing.com/search?format=rss&setlang=zh-Hans&mkt=zh-CN&count=30&q={quote_plus(query)}"
    response = session.get(url, timeout=(8, 18))
    response.raise_for_status()
    root = ElementTree.fromstring(response.content)
    hits = []
    for item in root.findall(".//item"):
        link = clean_text(item.findtext("link"))
        title = clean_text(item.findtext("title"))
        if link:
            hits.append(SearchHit(link, title, "Bing RSS", query))
    return hits


def discover_baidu(session: requests.Session, query: str) -> list[SearchHit]:
    url = f"https://www.baidu.com/s?rn=20&ie=utf-8&wd={quote_plus(query)}"
    response = session.get(url, timeout=(8, 18))
    response.raise_for_status()
    response.encoding = response.apparent_encoding or "utf-8"
    soup = BeautifulSoup(response.text, "lxml")
    hits = []
    for result in soup.select("div.result, div.c-container"):
        anchor = result.select_one("h3 a") or result.select_one("a")
        if not anchor or not anchor.get("href"):
            continue
        # 百度结果节点在 mu 属性中提供规范目标；不逐条请求跳转页，以控制运行时间。
        target = result.get("mu") or anchor.get("mu") or anchor["href"]
        title = anchor.get_text(" ", strip=True)
        hits.append(SearchHit(target, title, "百度", query))
    return hits


def discover_bing_html(session: requests.Session, query: str) -> list[SearchHit]:
    url = f"https://www.bing.com/search?setlang=zh-Hans&mkt=zh-CN&count=30&q={quote_plus(query)}"
    response = session.get(url, timeout=(8, 18))
    response.raise_for_status()
    response.encoding = response.apparent_encoding or "utf-8"
    soup = BeautifulSoup(response.text, "lxml")
    hits = []
    for item in soup.select("li.b_algo h2 a"):
        target = item.get("href", "")
        if target:
            hits.append(SearchHit(target, item.get_text(" ", strip=True), "Bing HTML", query))
    return hits


def load_seed_hits() -> list[SearchHit]:
    if not SEED_PATH.exists():
        return []
    hits = []
    for line in SEED_PATH.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if line and not line.startswith("#"):
            hits.append(SearchHit(line, provider="固定官方入口", query="seed"))
    return hits


def configured_official_indexes() -> list[str]:
    if not OFFICIAL_INDEX_PATH.exists():
        return []
    return [
        normalize_url(line)
        for line in OFFICIAL_INDEX_PATH.read_text(encoding="utf-8").splitlines()
        if line.strip() and not line.lstrip().startswith("#")
    ]


def official_index_urls(records: list[dict]) -> list[str]:
    """组合固定栏目与历史官方链接所属栏目，减少对搜索引擎收录速度的依赖。"""
    urls = set(configured_official_indexes())
    for record in records:
        url = normalize_url(record.get("官方原文链接", ""))
        if not is_official_url(url):
            continue
        parsed = urlparse(url)
        # 常见税务站路径：栏目/YYYYMM/tYYYYMMDD_xxx.html。
        index_path = re.sub(r"/20\d{4}/[^/]+$", "/", parsed.path)
        if index_path == parsed.path:
            continue
        urls.add(urlunparse((parsed.scheme, parsed.netloc, index_path, "", "", "")))
    return sorted(url for url in urls if url)


def date_from_listing_url(url: str) -> date | None:
    patterns = (
        r"/(20\d{2})(\d{2})/t(20\d{2})(\d{2})(\d{2})_",
        r"/(20\d{2})[-/](\d{1,2})[-/](\d{1,2})/",
    )
    for pattern in patterns:
        match = re.search(pattern, url)
        if not match:
            continue
        groups = match.groups()
        values = groups[-3:]
        try:
            return date(*map(int, values))
        except ValueError:
            continue
    return None


def discover_official_index(session: requests.Session, index_url: str) -> list[SearchHit]:
    """直接扫描官方公告栏目中的近期文章链接。"""
    response = session.get(index_url, timeout=(9, 28), allow_redirects=True)
    response.raise_for_status()
    response.encoding = response.apparent_encoding or response.encoding or "utf-8"
    soup = BeautifulSoup(response.text, "lxml")
    cutoff = TODAY - timedelta(days=OFFICIAL_INDEX_LOOKBACK_DAYS)
    hits: dict[str, SearchHit] = {}
    hints = (*DOC_TYPES, "税务文书送达公告", "公告送达", "稽查局")
    for anchor in soup.find_all("a", href=True):
        target = normalize_url(urljoin(response.url, anchor["href"]))
        if not is_official_url(target) or target == normalize_url(response.url):
            continue
        label = clean_text(anchor.get_text(" ", strip=True) or anchor.get("title", ""))
        context = clean_text(anchor.parent.get_text(" ", strip=True))[:240] if anchor.parent else label
        listing_date = date_from_listing_url(target)
        if listing_date and listing_date < cutoff:
            continue
        if not any(hint in f"{label} {context}" for hint in hints):
            continue
        hits[target] = SearchHit(target, label or context, "税务机关公告栏目", index_url)
        if len(hits) >= 120:
            break
    return list(hits.values())


def relevant_candidate(hit: SearchHit) -> bool:
    url = normalize_url(hit.url)
    if not url.startswith(("http://", "https://")):
        return False
    host = urlparse(url).netloc.lower()
    if any(token in host for token in ("bing.com", "baidu.com", "google.com", "sogou.com", "so.com")):
        return False
    title = clean_text(hit.title)
    return (
        not title
        or any(doc_type in title for doc_type in DOC_TYPES)
        or bool(re.search(r"税[一二三四]?稽[一二三四]?[处罚]", title))
        or is_official_url(url)
    )


def hit_priority(hit: SearchHit) -> int:
    """失败复核优先于官方栏目，官方栏目优先于搜索引擎。"""
    if hit.query == "retry-queue":
        return 40
    if hit.query == "pending-or-invalid":
        return 30
    if hit.provider in {"税务机关公告栏目", "固定官方入口"}:
        return 20
    return 10


def merge_search_hits(*hit_groups: Iterable[SearchHit]) -> list[SearchHit]:
    merged: dict[str, SearchHit] = {}
    for hit in (item for group in hit_groups for item in group):
        hit.url = normalize_url(hit.url)
        if not relevant_candidate(hit):
            continue
        current = merged.get(hit.url)
        if current is None or hit_priority(hit) > hit_priority(current):
            merged[hit.url] = hit
    return sorted(merged.values(), key=lambda item: (-hit_priority(item), item.url))


def discover_candidates(
    full_run: bool,
    errors: list[str],
    existing_records: list[dict],
    discovery_audits: list[dict] | None = None,
) -> list[SearchHit]:
    queries = SEARCH_QUERIES_FULL if full_run else SEARCH_QUERIES_RECENT
    hits = load_seed_hits()
    tasks: list[tuple] = []
    for index, query in enumerate(queries):
        for provider in (discover_bing_rss, discover_bing_html):
            tasks.append((provider, query))
        # 第二个独立搜索引擎；错开请求，避免给公开服务造成突发压力。
        if index % 2 == 0:
            tasks.append((discover_baidu, query))
    for index_url in official_index_urls(existing_records):
        tasks.append((discover_official_index, index_url))

    # 搜索入口彼此独立，可并行；每个任务使用自己的会话，避免跨线程共享连接状态。
    with ThreadPoolExecutor(max_workers=8) as executor:
        future_map = {
            executor.submit(provider, new_session(), query): (provider, query)
            for provider, query in tasks
        }
        for future in as_completed(future_map):
            provider, query = future_map[future]
            try:
                hits.extend(future.result())
            except Exception as exc:
                errors.append(f"{provider.__name__}: {query}: {type(exc).__name__}: {exc}")
                if discovery_audits is not None and is_official_url(query):
                    discovery_audits.append({
                        "checked_at": RUN_NOW.isoformat(timespec="seconds"),
                        "provider": "税务机关公告栏目",
                        "query": query,
                        "url": normalize_url(query),
                        "final_url": normalize_url(query),
                        "status_code": None,
                        "page_state": "暂时无法访问",
                        "title": "",
                        "error": f"{type(exc).__name__}: {exc}",
                        "records": 0,
                    })

    return merge_search_hits(hits)


def historical_revalidation_hits(records: list[dict]) -> list[SearchHit]:
    """每日重新核验历史待核验及失效链接，历史记录本身永不删除。"""
    hits: dict[str, SearchHit] = {}
    for record in records:
        suspicious_subject = suspicious_party_name(record.get("当事人名称"))
        unresolved_conflict = "存在冲突" in clean_text(record.get("备注"))
        needs_check = (
            clean_text(record.get("核验状态")) == "待核验"
            or clean_text(record.get("页面当前状态")) not in NORMAL_STATES
            or suspicious_subject
            or unresolved_conflict
        )
        if not needs_check:
            continue
        fields = (
            ("官方原文链接",)
            if (suspicious_subject or unresolved_conflict)
            and clean_text(record.get("页面当前状态")) in NORMAL_STATES
            else ("官方原文链接", "附件链接", "备用来源链接")
        )
        for field_name in fields:
            for raw_url in clean_text(record.get(field_name)).split(";"):
                url = normalize_url(raw_url)
                if url.startswith(("http://", "https://")) and url not in hits:
                    hits[url] = SearchHit(
                        url=url,
                        title=clean_text(record.get("页面标题")),
                        provider="历史记录复核",
                        query="pending-or-invalid",
                    )
    return list(hits.values())


def detect_page_state(status_code: int | None, text: str, error: str) -> str:
    if status_code in (404, 410):
        return "页面已删除"
    if status_code in (403, 408, 412, 425, 429) or status_code is None or status_code >= 500 or error:
        return "暂时无法访问"
    lowered = text.lower()
    if any(doc_type in text for doc_type in DOC_TYPES):
        return "正常"
    if any(token in lowered for token in ("页面不存在", "您访问的页面不存在", "404 not found", "content not found")):
        return "页面已删除"
    return "内容不匹配"


def extract_binary_text(content: bytes, content_type: str, url: str) -> str:
    path = urlparse(url).path.lower()
    try:
        if content.startswith(b"%PDF") or "pdf" in content_type or path.endswith(".pdf"):
            reader = PdfReader(io.BytesIO(content))
            return "\n".join((page.extract_text() or "") for page in reader.pages[:120])
        if (
            "wordprocessingml" in content_type
            or path.endswith(".docx")
            or content.startswith(b"PK\x03\x04") and b"word/" in content[:5000]
        ):
            document = Document(io.BytesIO(content))
            parts = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    parts.append("\t".join(cell.text for cell in row.cells))
            return "\n".join(parts)
    except Exception:
        return ""
    return ""


def valid_attachment_content(content: bytes) -> bool:
    """只接受实际 PDF/Word 文件，拒绝 HTTP 200 的防盗链提示页。"""
    if content.startswith(b"%PDF") or content.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"):
        return True
    if not content.startswith(b"PK\x03\x04"):
        return False
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            return any(name.startswith("word/") for name in archive.namelist())
    except zipfile.BadZipFile:
        return False


def fetch_attachment(session: requests.Session, url: str, label: str, referer: str = "") -> dict:
    item = {"url": normalize_url(url), "label": clean_text(label), "ok": False, "text": "", "state": "暂时无法访问"}
    try:
        headers = {"Referer": referer} if referer else None
        response = session.get(url, headers=headers, timeout=(8, 28), allow_redirects=True)
        content = response.content[:25 * 1024 * 1024]
        content_type = response.headers.get("Content-Type", "").split(";")[0].lower()
        attachment_ok = response.ok and valid_attachment_content(content)
        item.update({
            "url": normalize_url(response.url),
            "ok": attachment_ok,
            "text": extract_binary_text(content, content_type, response.url) if attachment_ok else "",
            "state": "附件正常" if attachment_ok else detect_page_state(response.status_code, "", ""),
            "status_code": response.status_code,
        })
        if response.ok and not attachment_ok:
            item["state"] = "需要人工核验"
            item["error"] = "附件地址返回的不是可验证的 PDF/Word 文件"
    except requests.RequestException as exc:
        item["error"] = f"{type(exc).__name__}: {exc}"
    return item


def fetch_candidate(hit: SearchHit) -> FetchResult:
    session = new_session()
    try:
        response = session.get(hit.url, timeout=(9, 28), allow_redirects=True)
        content = response.content[:30 * 1024 * 1024]
        content_type = response.headers.get("Content-Type", "").split(";")[0].lower()
        final_url = normalize_url(response.url)
        is_html = "html" in content_type or content.lstrip().startswith((b"<!DOCTYPE", b"<html", b"<HTML"))
        attachments: list[dict] = []
        title = hit.title
        if is_html:
            response.encoding = response.apparent_encoding or response.encoding or "utf-8"
            soup = BeautifulSoup(response.text, "lxml")
            for tag in soup(["script", "style", "noscript", "svg"]):
                tag.decompose()
            title_tag = soup.find("title")
            if title_tag:
                title = clean_text(title_tag.get_text(" ", strip=True))
            body = soup.get_text("\n", strip=True)
            attachment_links: dict[str, str] = {}
            for anchor in soup.find_all("a", href=True):
                label = clean_text(anchor.get_text(" ", strip=True))
                href = urljoin(response.url, anchor["href"])
                lower = href.lower()
                if (
                    any(ext in lower for ext in (".pdf", ".doc", ".docx", ".wps", ".zip", "download", "downfile"))
                    or any(doc_type in label for doc_type in DOC_TYPES)
                    or "附件" in label
                ):
                    if href.startswith(("http://", "https://")):
                        attachment_links[normalize_url(href)] = label
            # 只核验与目标文书相关的前 12 个附件，防止栏目页包含大量无关下载。
            selected = list(attachment_links.items())[:12]
            with ThreadPoolExecutor(max_workers=4) as executor:
                futures = [
                    executor.submit(fetch_attachment, session, url, label, response.url)
                    for url, label in selected
                ]
                for future in as_completed(futures):
                    attachments.append(future.result())
            text = body
        else:
            text = extract_binary_text(content, content_type, final_url)
            if not title:
                title = Path(urlparse(final_url).path).name
        state = detect_page_state(response.status_code, text, "")
        if any(att.get("ok") for att in attachments):
            state = "附件正常"
        return FetchResult(
            hit.url, final_url, response.ok, response.status_code, content_type, content,
            compact_multiline(text), clean_text(title), state, attachments=attachments,
        )
    except requests.RequestException as exc:
        return FetchResult(
            hit.url, hit.url, False, None, "", b"", "", clean_text(hit.title),
            "暂时无法访问", f"{type(exc).__name__}: {exc}",
        )
    except Exception as exc:
        return FetchResult(
            hit.url, hit.url, False, None, "", b"", "", clean_text(hit.title),
            "需要人工核验", f"{type(exc).__name__}: {exc}",
        )


def extract_publication_date(text: str, url: str) -> str:
    # 优先“发布时间/发布日期/时间”，避免误取正文中的检查期间或决定日期。
    patterns = [
        r"(?:发布时间|发布日期|发布于|时间|信息发布时间)\s*[:：]?\s*" + DATE_PATTERN.pattern,
        r"(?:发布时间|发布日期)\D{0,20}(20\d{2})-(\d{1,2})-(\d{1,2})",
    ]
    for pattern in patterns:
        match = re.search(pattern, text[:6000], re.I)
        if match:
            groups = match.groups()[-3:]
            try:
                return date(*map(int, groups)).isoformat()
            except ValueError:
                pass
    # 税务网站 URL 通常包含 YYYYMM 或 YYYY/MM/DD，属于可审计的发布日期线索。
    for pattern in (
        r"(?:/|t)(20\d{2})[/-]?(\d{2})[/-]?(\d{2})",
        r"P0(20\d{2})(\d{2})(\d{2})",
    ):
        match = re.search(pattern, url, re.I)
        if match:
            try:
                return date(*map(int, match.groups())).isoformat()
            except ValueError:
                pass
    return ""


def locate_region(title: str, text: str, url: str) -> tuple[str, str, str]:
    haystack = f"{title}\n{text[:5000]}\n{url}"
    province = ""
    host_prefix = urlparse(url).netloc.lower().split(".")[0]
    if host_prefix in DOMAIN_PROVINCES:
        province = DOMAIN_PROVINCES[host_prefix]
    for full in PROVINCES:
        if not province and full in haystack:
            province = full
            break
    if not province:
        for alias, full in PROVINCE_ALIASES.items():
            if alias in haystack:
                province = full
                break

    agency_text = haystack.replace("国家税务总局", " ")
    if province:
        for token in (province, province.replace("省", "").replace("市", "").replace("自治区", "")):
            agency_text = agency_text.replace(token, "", 1)
    city = ""
    match = re.search(r"([\u4e00-\u9fff]{2,10}(?:市|自治州|地区|盟))", agency_text)
    if match:
        city = match.group(1)
    if province in {"北京市", "天津市", "上海市", "重庆市"}:
        city = province

    county = ""
    county_text = haystack.replace("国家税务总局", " ").replace("税务总局", " ")
    for token in (province, city):
        if token:
            county_text = county_text.replace(token, " ")
    for match in re.finditer(r"([\u4e00-\u9fff]{2,10}(?:区|县|旗))税务局", county_text):
        value = match.group(1)
        value = re.split(r"[到在向由至请内]", value)[-1]
        if 2 <= len(value) <= 8 and "自治区" not in value and "税务" not in value:
            county = value
            break
    return province, city, county


def extract_agency(title: str, text: str) -> tuple[str, str]:
    candidates = re.findall(
        r"国家税务总局[\u4e00-\u9fff·]{2,45}?(?:税务局(?:(?:第[一二三四五])?稽查局)?|稽查局)",
        f"{title}\n{text[-4000:]}",
    )
    cleaned = []
    for candidate in candidates:
        candidate = re.split(r"(?:关于|公告|送达)", candidate)[0]
        if candidate not in cleaned:
            cleaned.append(candidate)
    inspection = next((item for item in cleaned if "稽查局" in item), "")
    agency = inspection or (cleaned[0] if cleaned else "")
    if not inspection:
        match = re.search(r"([\u4e00-\u9fff]{2,25}(?:第一|第二|第三|第四)?稽查局)", f"{title}\n{text}")
        inspection = match.group(1) if match else ""
    return clean_text(agency), clean_text(inspection)


def candidate_subjects(title: str, text: str) -> list[str]:
    names = []
    rejected = {
        "来源", "字号", "访问次数", "本站热词", "联系电话", "联系人", "下载", "附件",
        "纳税人识别号", "统一社会信用代码", "社会信用代码", "发布时间", "发布日期",
        "国家税务总局", "税务局", "稽查局", "正文下载", "打印本页", "信息公开",
    }
    patterns = [
        # 公告常以独立的“公司名称：”称呼当事人，并不总是公布信用代码。
        # 限定行首与企业后缀，避免把案情中提到的其他公司当成受送达人。
        r"(?:^|\n)\s*([\u4e00-\u9fffA-Za-z0-9·（）()]{2,70}?"
        r"(?:有限责任公司|股份有限公司|有限公司|公司|合作社|事务所|经营部|商行|商贸部|厂|店|中心))\s*[:：]",
        r"(?:^|\n|\s)\d{0,3}\s*([\u4e00-\u9fffA-Za-z0-9·（）()]{2,70}?"
        r"(?:有限责任公司|股份有限公司|有限公司|公司|合作社|事务所|经营部|商行|商贸部|厂|店|中心))"
        r"\s+(?:[0-9A-HJ-NPQRTUWXY]{18}|\d{15})\s+[^\n。；;]{0,35}?税[^\n。；;]{0,12}?[处罚]",
        r"(?:^|\n|[。；;])\s*(?:现向|向|对)?\s*([^\n。；;：:]{2,70}?"
        r"(?:有限责任公司|股份有限公司|有限公司|公司|合作社|事务所|经营部|商行|商贸部|厂|店|中心))"
        r"\s*[:：]?\s*[（(](?:统一社会信用代码|纳税人识别号|社会信用代码)",
        r"^([^\n。；;：:]{2,70})\s*[:：]?\s*[（(](?:统一社会信用代码|纳税人识别号|社会信用代码)",
        r"关于送达([\u4e00-\u9fffA-Za-z0-9·（）()]{2,70}?"
        r"(?:有限责任公司|股份有限公司|有限公司|公司|合作社|事务所|经营部|商行|厂|店|中心))",
        r"关于(?:对|送达)?([\u4e00-\u9fff·]{2,12})的?税务文书",
    ]
    haystack = f"{title}\n{text}"
    for pattern in patterns:
        for match in re.finditer(pattern, haystack, re.M):
            name = clean_text(match.group(1)).strip("：:，,。 ")
            name = re.sub(r"^(?:公告|附件|国家税务总局|现向|向)", "", name)
            if "对" in name and name.rsplit("对", 1)[-1].endswith(("公司", "厂", "店", "中心", "合作社", "事务所", "经营部", "商行", "商贸部")):
                name = name.rsplit("对", 1)[-1]
            name = re.sub(r"^.*(?:年\d{1,2}月\d{1,2}日|经查|我局)", "", name)
            # 个别官网正文会把企业后缀重复成“有限公司公司”，应归一为同一当事人，
            # 否则同页的处理、处罚文号会被错误拆到两个主体块中。
            name = re.sub(r"(有限责任公司|股份有限公司|有限公司|公司)公司$", r"\1", name)
            name = name.rstrip("的").strip()
            if (
                2 <= len(name) <= 70
                and name not in names
                and name not in rejected
                and not any(doc in name for doc in DOC_TYPES)
                and not any(token in name for token in ("发布时间", "访问次数", "本站热词", "纳税人识别号"))
                and not suspicious_party_name(name)
            ):
                names.append(name)
    return names[:30]


def subject_block(text: str, subject: str, all_subjects: list[str]) -> str:
    positions = [match.start() for match in re.finditer(re.escape(subject), text)]
    if not positions:
        return text
    position = positions[0]
    best_score = -1
    for candidate in positions:
        window = text[candidate:candidate + 700]
        score = 0
        if DOC_NO_PATTERN.search(window):
            score += 2
        if re.search(r"(?:[0-9A-HJ-NPQRTUWXY]{18}|\d{15})", window, re.I):
            score += 1
        if score >= best_score:
            position = candidate
            best_score = score
    end = len(text)
    for other in all_subjects:
        if other == subject:
            continue
        other_pos = text.find(other, position + len(subject))
        if other_pos >= 0:
            end = min(end, other_pos)
    return text[position:end]


def extract_doc_numbers(text: str) -> list[tuple[str, str]]:
    found: list[tuple[str, str]] = []
    for match in DOC_NO_PATTERN.finditer(text):
        number = normalized_doc_no(match.group("no"))
        if "不罚" in number or "不予" in number:
            continue
        context = text[max(0, match.start() - 35):match.end() + 15]
        doc_type = doc_type_from_number(number, context)
        item = (doc_type, number)
        if item not in found:
            found.append(item)
    return found


def document_number_year(value: object) -> int | None:
    match = re.search(r"〔(20\d{2})〕", normalized_doc_no(value))
    return int(match.group(1)) if match else None


def excluded_only_title(value: object) -> bool:
    title = clean_text(value)
    return bool(
        any(marker in title for marker in EXCLUDED_ONLY_TITLE_MARKERS)
        and not any(doc_type in title for doc_type in DOC_TYPES)
    )


def provisional_out_of_scope(record: dict) -> bool:
    """回滚本次运行中无法证明属于 2026 年公开范围的旧年度文书。"""
    if excluded_only_title(record.get("页面标题")):
        return True
    published = iso_date(record.get("官方发布日期")) or iso_date(record.get("第三方收录日期"))
    year = document_number_year(record.get("决定书文号"))
    return bool(
        not published
        and year is not None
        and year < 2026
        and iso_date(record.get("首次发现日期")) == TODAY.isoformat()
    )


def find_uscc(block: str, subject: str) -> str:
    if not subject.endswith(ENTITY_SUFFIXES):
        return ""
    match = re.search(
        r"(?:统一社会信用代码|纳税人识别号|社会信用代码)\s*[:：]?\s*"
        r"([0-9A-HJ-NPQRTUWXY]{18})",
        block[:1600],
        re.I,
    )
    if not match:
        match = re.search(
            r"^\s*" + re.escape(subject) + r"\s+([0-9A-HJ-NPQRTUWXY]{18})(?:\s|$)",
            block[:600],
            re.I,
        )
    return match.group(1).upper() if match else ""


def find_legal_rep(block: str) -> str:
    match = re.search(r"法定代表人\s*[:：为]?\s*([\u4e00-\u9fff·]{2,12})", block[:3000])
    if not match:
        return ""
    value = match.group(1)
    if any(token in value for token in INVALID_LEGAL_REP_TOKENS):
        return ""
    return value


def extract_amounts(text: str) -> dict[str, float | None]:
    compact = clean_text(text)
    values: dict[str, float | None] = {
        "追缴税款金额": None, "滞纳金金额": None, "罚款金额": None, "没收违法所得金额": None,
    }
    tax_pairs: list[tuple[str, float]] = []
    aggregate_sums: list[float] = []
    for sentence in re.split(r"[。；;\n]", compact):
        action = re.search(r"(?:应|决定|需要|合计需|本次检查决定)?(?:追缴|补缴)", sentence)
        if not action:
            continue
        segment = sentence[action.start():]
        sentence_pairs: list[tuple[str, float]] = []
        found_named = False
        for tax in TAX_NAMES:
            for match in re.finditer(re.escape(tax) + r"[^0-9]{0,15}" + AMOUNT_PATTERN, segment):
                amount = parse_money(match.group(1))
                if amount is not None:
                    sentence_pairs.append((tax, amount))
                    found_named = True
        if not found_named:
            match = re.search(r"(?:追缴|补缴)[^。；\n]{0,50}?" + AMOUNT_PATTERN, segment)
            if match:
                amount = parse_money(match.group(1))
                if amount is not None:
                    sentence_pairs.append(("税款", amount))
        if sentence_pairs:
            unique_sentence_pairs = list(dict.fromkeys(sentence_pairs))
            tax_pairs.extend(unique_sentence_pairs)
            if any(token in sentence for token in ("合计", "综上", "共应", "本次检查决定")):
                aggregate_sums.append(sum(amount for _, amount in unique_sentence_pairs))
    if aggregate_sums:
        values["追缴税款金额"] = round(max(aggregate_sums), 2)
    elif tax_pairs:
        unique_pairs = list(dict.fromkeys(tax_pairs))
        values["追缴税款金额"] = round(sum(amount for _, amount in unique_pairs), 2)

    keyword_patterns = {
        "滞纳金金额": (r"滞纳金(?:合计|金额)?\s*(?:为|人民币|[:：])?\s*" + AMOUNT_PATTERN,),
        "罚款金额": (
            r"(?:处以|处罚|处)?罚款[^。；\n\d]{0,24}" + AMOUNT_PATTERN,
            r"(?:处以|处罚|处)\s*" + AMOUNT_PATTERN + r"\s*的?罚款",
        ),
        "没收违法所得金额": (r"没收(?:违法所得)?[^。；\n\d]{0,24}" + AMOUNT_PATTERN,),
    }
    for field_name, patterns in keyword_patterns.items():
        amounts = [
            parse_money(match.group(1))
            for pattern in patterns
            for match in re.finditer(pattern, compact)
        ]
        amounts = [value for value in amounts if value is not None]
        if amounts:
            values[field_name] = round(max(amounts), 2)
    return values


def extract_tax_types(text: str) -> str:
    return "、".join(tax for tax in TAX_NAMES if tax in text)


def extract_fact(text: str) -> str:
    sentences = [clean_text(s) for s in re.split(r"(?<=[。；;])|\n", text)]
    chosen = []
    for sentence in sentences:
        if any(token in sentence for token in ("经查", "违法事实", "虚开", "骗取", "偷税", "少缴", "未申报")):
            if not any(token in sentence for token in ("搜索", "栏目", "当前位置")) and len(sentence) >= 18:
                chosen.append(sentence)
        if sum(map(len, chosen)) >= 700:
            break
    return clean_text(" ".join(chosen))[:1000]


def extract_result(text: str) -> str:
    sentences = [clean_text(s) for s in re.split(r"(?<=[。；;])|\n", text)]
    chosen = []
    for sentence in sentences:
        if any(token in sentence for token in ("决定追缴", "应追缴", "补缴", "处以罚款", "处罚款", "没收违法所得", "予以公告送达", "撤销上述")):
            if len(sentence) >= 12:
                chosen.append(sentence)
        if sum(map(len, chosen)) >= 700:
            break
    return clean_text(" ".join(chosen))[:1000]


def attachment_for_doc(attachments: list[dict], doc_type: str, doc_no: str) -> dict | None:
    normalized = normalized_doc_no(doc_no)
    for attachment in attachments:
        haystack = f"{attachment.get('label', '')}\n{attachment.get('text', '')}"
        if normalized and normalized in normalized_doc_no(haystack):
            return attachment
    for attachment in attachments:
        if doc_type in f"{attachment.get('label', '')}\n{attachment.get('text', '')}":
            return attachment
    return next((attachment for attachment in attachments if attachment.get("ok")), None)


def extract_decision_date(attachment_text: str, doc_no: str, publication_date: str) -> str:
    if not attachment_text:
        return ""
    dates = []
    for match in DATE_PATTERN.finditer(attachment_text):
        try:
            value = date(*map(int, match.groups()))
        except ValueError:
            continue
        if 2020 <= value.year <= TODAY.year:
            dates.append(value)
    if not dates:
        return ""
    publication = date.fromisoformat(publication_date) if publication_date else TODAY
    eligible = [value for value in dates if value <= publication]
    # 文书落款通常是正文中最后一个、不晚于发布日期的日期。
    return max(eligible or dates).isoformat()


def parse_fetch(hit: SearchHit, fetched: FetchResult, start_date: date, full_run: bool) -> tuple[list[dict], dict]:
    audit = {
        "checked_at": RUN_NOW.isoformat(timespec="seconds"),
        "provider": hit.provider,
        "query": hit.query,
        "url": hit.url,
        "final_url": fetched.final_url,
        "status_code": fetched.status_code,
        "page_state": fetched.state,
        "title": fetched.title,
        "error": fetched.error,
        "records": 0,
    }
    if not fetched.ok or not fetched.text:
        return [], audit

    attachment_corpus = "\n".join(
        f"{attachment.get('label', '')}\n{attachment.get('text', '')}"
        for attachment in fetched.attachments
    )
    title_and_text = f"{fetched.title}\n{fetched.text}\n{attachment_corpus}"
    publication_date = extract_publication_date(title_and_text, fetched.final_url)
    if publication_date:
        published = date.fromisoformat(publication_date)
        if published < date(2026, 1, 1):
            audit["skip_reason"] = "发布日期早于 2026-01-01"
            return [], audit
        is_historical_revalidation = hit.query in {"pending-or-invalid", "retry-queue"}
        if not full_run and not is_historical_revalidation and published < start_date:
            audit["skip_reason"] = "不在本次增量时间窗"
            return [], audit

    # 明确只有排除文书、且正文没有目标决定书/文号时不收录。
    excluded = [name for name in EXCLUDED_TITLES if name in fetched.title]
    doc_numbers_global = extract_doc_numbers(title_and_text)
    if excluded or excluded_only_title(fetched.title):
        if excluded_only_title(fetched.title):
            excluded.append("税务事项通知书（标题未公开目标决定书）")
        audit["skip_reason"] = "排除文书类型：" + "、".join(excluded)
        return [], audit
    if not any(doc_type in title_and_text for doc_type in DOC_TYPES) or not doc_numbers_global:
        audit["skip_reason"] = "未同时确认目标文书类型和决定书文号"
        return [], audit
    if not publication_date and all(
        (document_number_year(doc_no) or 0) < 2026 for _, doc_no in doc_numbers_global
    ):
        audit["skip_reason"] = "无法确认2026年后发布日期，且文号年份早于2026"
        return [], audit

    subjects = candidate_subjects(fetched.title, f"{fetched.text}\n{attachment_corpus}")
    if not subjects:
        audit["skip_reason"] = "无法确认当事人"
        return [], audit

    province, city, county = locate_region(fetched.title, fetched.text, fetched.final_url)
    agency, inspection = extract_agency(fetched.title, fetched.text)
    level = source_level(fetched.final_url)
    records: list[dict] = []

    for subject in subjects:
        block = subject_block(f"{fetched.text}\n{attachment_corpus}", subject, subjects)
        doc_numbers = extract_doc_numbers(block)
        if len(subjects) == 1:
            # 单一当事人页面常把处理、处罚文号分散在不同正文段落；主体不存在歧义时，
            # 合并全页已确认文号，避免“最佳段落”只命中其中一份文书。
            for item in doc_numbers_global:
                if item not in doc_numbers:
                    doc_numbers.append(item)
        if not doc_numbers:
            continue
        uscc = find_uscc(block, subject)
        legal_rep = find_legal_rep(block)

        for doc_type, doc_no in doc_numbers:
            attachment = attachment_for_doc(fetched.attachments, doc_type, doc_no)
            attachment_text = compact_multiline(attachment.get("text", "")) if attachment else ""
            # 附件明确对应当前文号时，以该附件为事实与金额证据，避免同页两份文书相互串值。
            evidence = attachment_text or block
            amounts = extract_amounts(evidence)
            fact = extract_fact(evidence)
            result = extract_result(evidence)

            # 金额只分配给语义对应的文书，避免同一公告中的税款和罚款同时复制到两行。
            if doc_type == "税务处理决定书":
                amounts["罚款金额"] = None
                amounts["没收违法所得金额"] = None
            else:
                amounts["追缴税款金额"] = None
                amounts["滞纳金金额"] = None

            substantive_result = bool(
                result and any(token in result for token in ("追缴", "补缴", "罚款", "没收违法所得"))
            )
            has_detail = bool(fact or substantive_result or any(value is not None for value in amounts.values()))
            attachment_url = attachment.get("url", "") if attachment else ""
            attachment_ok = bool(attachment and attachment.get("ok"))
            completeness = "完整文书" if has_detail else "仅公告送达/仅文号线索"
            verified = "已核验" if level != "第三方待核验" and fetched.state in NORMAL_STATES else "待核验"
            page_state = fetched.state
            if attachment_ok:
                page_state = "附件正常"

            official_link = fetched.final_url if level != "第三方待核验" else ""
            backup_link = "" if level != "第三方待核验" else fetched.final_url
            notes = []
            if attachment_ok and not has_detail:
                notes.append("官方附件可下载，但尚未提取到可核实的案情或金额，按文号线索收录。")
            if "撤销" in fetched.title or "撤销上述" in fetched.text:
                notes.append("官方页面显示该决定书已被撤销。")
            if "告知书" in fetched.title and doc_type not in fetched.title:
                notes.append("页面标题与正文/附件文书类型不一致，需人工复核。")
                verified = "待核验"
                page_state = "内容不匹配"
            if not publication_date:
                notes.append("页面未提取到可确认的官方发布日期。")
                verified = "待核验"
            if level == "第三方待核验":
                notes.append("尚未找到官方原文；第三方来源仅作待核验线索。")

            record = {field_name: "" for field_name in FIELDS}
            record.update({
                "省份": province,
                "城市": city,
                "区县": county,
                "发布机关": agency,
                "稽查机构": inspection,
                "当事人名称": subject,
                "统一社会信用代码": uscc,
                "法定代表人": legal_rep,
                "文书类型": doc_type,
                "决定书文号": doc_no,
                "关联处理决定书文号": doc_no if doc_type == "税务处理决定书" else "",
                "关联处罚决定书文号": doc_no if doc_type == "税务行政处罚决定书" else "",
                "主要违法事实": fact,
                "涉及税种": extract_tax_types(evidence),
                "追缴税款金额": amounts["追缴税款金额"],
                "滞纳金金额": amounts["滞纳金金额"],
                "罚款金额": amounts["罚款金额"],
                "没收违法所得金额": amounts["没收违法所得金额"],
                "处理或处罚结果": result,
                "决定书作出日期": extract_decision_date(attachment_text, doc_no, publication_date),
                "官方发布日期": publication_date if level != "第三方待核验" else "",
                "第三方收录日期": publication_date if level == "第三方待核验" else "",
                "首次发现日期": TODAY.isoformat(),
                "最后核验日期": TODAY.isoformat(),
                "公开完整度": completeness,
                "来源级别": level,
                "核验状态": verified,
                "官方原文链接": official_link,
                "附件链接": attachment_url,
                "备用来源链接": backup_link,
                "页面标题": fetched.title,
                "页面当前状态": page_state,
                "备注": " ".join(notes),
            })
            record["文书唯一ID"] = build_unique_id(record)
            record["案件组ID"] = build_group_id(record)
            records.append(record)

    # 同一页面内按唯一 ID 去重。
    unique = {}
    for record in records:
        unique[record["文书唯一ID"]] = record
    records = list(unique.values())

    # 将同一当事人/机关/年度的两类文书互相写入关联文号。
    groups: dict[str, list[dict]] = {}
    for record in records:
        groups.setdefault(record["案件组ID"], []).append(record)
    for group_records in groups.values():
        processing = next((r["决定书文号"] for r in group_records if r["文书类型"] == "税务处理决定书"), "")
        penalty = next((r["决定书文号"] for r in group_records if r["文书类型"] == "税务行政处罚决定书"), "")
        for record in group_records:
            if processing:
                record["关联处理决定书文号"] = processing
            if penalty:
                record["关联处罚决定书文号"] = penalty

    audit["records"] = len(records)
    return records, audit


def read_csv_records(path: Path) -> list[dict]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle)
        records = []
        for row in reader:
            record = {field_name: row.get(field_name, "") for field_name in FIELDS}
            for field_name in MONEY_FIELDS:
                record[field_name] = parse_money(record[field_name])
            records.append(record)
        return records


def read_xlsx_records(path: Path) -> tuple[list[dict], list[dict]]:
    if not path.exists():
        return [], []
    workbook = load_workbook(path, data_only=False, read_only=False)
    records: list[dict] = []
    logs: list[dict] = []
    if "文书汇总" in workbook.sheetnames:
        sheet = workbook["文书汇总"]
        headers = [clean_text(cell.value) for cell in sheet[1]]
        for row in sheet.iter_rows(min_row=2):
            if not any(cell.value not in (None, "") for cell in row):
                continue
            item = {}
            for index, header in enumerate(headers):
                if header not in FIELDS:
                    continue
                cell = row[index]
                value = cell.value
                if header in LINK_FIELDS and cell.hyperlink:
                    value = cell.hyperlink.target
                if header in MONEY_FIELDS:
                    value = parse_money(value)
                elif header in DATE_FIELDS:
                    value = iso_date(value)
                item[header] = value if value is not None else ""
            records.append({field_name: item.get(field_name, "") for field_name in FIELDS})
    if "每日运行日志" in workbook.sheetnames:
        sheet = workbook["每日运行日志"]
        headers = [clean_text(cell.value) for cell in sheet[1]]
        for row in sheet.iter_rows(min_row=2, values_only=True):
            if not any(value not in (None, "") for value in row):
                continue
            logs.append({
                header: (row[index] if index < len(row) and row[index] is not None else "")
                for index, header in enumerate(headers) if header in LOG_FIELDS
            })
    workbook.close()
    return records, logs


def canonical_keys(record: dict) -> list[tuple]:
    doc_type = clean_text(record.get("文书类型"))
    doc_no = normalized_doc_no(record.get("决定书文号", ""))
    party = clean_text(record.get("当事人名称"))
    decision_date = iso_date(record.get("决定书作出日期"))
    url = normalize_url(record.get("官方原文链接") or record.get("备用来源链接") or "")
    keys = []
    if doc_type and doc_no:
        keys.append(("type_no", doc_type, doc_no))
    if party and doc_no:
        keys.append(("party_no", party, doc_no))
    if party and doc_type and decision_date:
        keys.append(("party_type_date", party, doc_type, decision_date))
    if url:
        keys.append(("url", url, doc_type, party))
    return keys


def comparison_text(value: object) -> str:
    return re.sub(r"\s+", " ", clean_text(value)).strip()


def suspicious_party_name(value: object) -> bool:
    name = comparison_text(value)
    return bool(name and any(marker in name for marker in INVALID_PARTY_MARKERS))


def append_note(record: dict, note: str) -> bool:
    note = clean_text(note)
    current = clean_text(record.get("备注"))
    if note and comparison_text(note) not in comparison_text(current):
        record["备注"] = f"{current} {note}".strip()
        return True
    return False


def informative_value(value: object) -> bool:
    return value is not None and clean_text(value) not in ("", "None", "nan")


def merge_record(existing: dict, incoming: dict) -> tuple[dict, bool]:
    merged = dict(existing)
    changed = False
    existing_priority = SOURCE_PRIORITY.get(clean_text(existing.get("来源级别")), 0)
    incoming_priority = SOURCE_PRIORITY.get(clean_text(incoming.get("来源级别")), 0)

    preserve = {"序号", "首次发现日期", "文书唯一ID", "案件组ID"}
    always_refresh = {"最后核验日期", "页面当前状态", "核验状态"}
    for field_name in FIELDS:
        old = merged.get(field_name, "")
        new = incoming.get(field_name, "")
        if field_name in preserve:
            if not informative_value(old) and informative_value(new):
                merged[field_name] = new
                changed = True
            continue
        if field_name in always_refresh and informative_value(new):
            if clean_text(old) != clean_text(new):
                merged[field_name] = new
                changed = True
            continue
        if not informative_value(new):
            continue
        if not informative_value(old):
            merged[field_name] = new
            changed = True
            continue
        comparable_old = money_string(old) if field_name in MONEY_FIELDS else comparison_text(old)
        comparable_new = money_string(new) if field_name in MONEY_FIELDS else comparison_text(new)
        if comparable_old == comparable_new:
            continue

        if (
            field_name == "当事人名称"
            and suspicious_party_name(old)
            and not suspicious_party_name(new)
        ):
            append_note(merged, f"当事人名称经官方原文复核，由“{old}”更正为“{new}”。")
            merged[field_name] = new
            changed = True
        elif field_name == "来源级别" and incoming_priority > existing_priority:
            merged[field_name] = new
            changed = True
        elif field_name == "官方原文链接" and incoming_priority > existing_priority:
            prior_url = normalize_url(old)
            merged[field_name] = new
            if prior_url and not merged.get("备用来源链接"):
                merged["备用来源链接"] = prior_url
            changed = True
        elif field_name == "备用来源链接":
            urls = [normalize_url(value) for value in f"{old};{new}".split(";") if clean_text(value)]
            combined = ";".join(dict.fromkeys(urls))
            if combined != old:
                merged[field_name] = combined
                changed = True
        elif field_name == "公开完整度":
            if old != "完整文书" and new == "完整文书":
                merged[field_name] = new
                changed = True
        elif field_name in {"主要违法事实", "处理或处罚结果"}:
            attachment_verified = (
                clean_text(incoming.get("来源级别")) == "税务机关官网"
                and clean_text(incoming.get("页面当前状态")) == "附件正常"
            )
            if attachment_verified or len(clean_text(new)) > len(clean_text(old)):
                append_note(merged, f"字段“{field_name}”已根据官方附件重新解析。")
                merged[field_name] = new
                changed = True
        elif incoming_priority > existing_priority and field_name in {
            "省份", "城市", "区县", "发布机关", "稽查机构", "统一社会信用代码",
            "法定代表人", "决定书作出日期", "官方发布日期", "附件链接", "页面标题",
        }:
            append_note(merged, f"字段“{field_name}”存在来源冲突：原值“{old}”，采用更权威来源值“{new}”。")
            merged[field_name] = new
            changed = True
        elif field_name not in {"备注", "最后核验日期", "页面当前状态", "核验状态"}:
            changed = append_note(
                merged,
                f"字段“{field_name}”存在冲突：保留原值“{old}”，新值“{new}”未静默覆盖。",
            ) or changed

    merged["首次发现日期"] = iso_date(existing.get("首次发现日期")) or iso_date(incoming.get("首次发现日期")) or TODAY.isoformat()
    # 合并两个历史副本不等于重新访问官网，只有真实抓取结果才能带来新核验日期。
    merged["最后核验日期"] = max(iso_date(existing.get("最后核验日期")), iso_date(incoming.get("最后核验日期")))
    merged["文书唯一ID"] = clean_text(existing.get("文书唯一ID")) or build_unique_id(merged)
    merged["案件组ID"] = clean_text(existing.get("案件组ID")) or build_group_id(merged)
    return merged, changed


def sanitize_record(record: dict) -> None:
    """清除无法由明确标签支持的地域/主体字段，不用上下文猜测补值。"""
    for field_name in ("官方原文链接", "附件链接", "备用来源链接"):
        links = [normalize_url(item) for item in clean_text(record.get(field_name)).split(";") if clean_text(item)]
        record[field_name] = ";".join(dict.fromkeys(link for link in links if link))

    title = clean_text(record.get("页面标题"))
    agency = clean_text(record.get("发布机关"))
    inspection = clean_text(record.get("稽查机构"))
    url = clean_text(record.get("官方原文链接") or record.get("备用来源链接"))
    province, city, county = locate_region(title, f"{agency}\n{inspection}", url)
    if province:
        record["省份"] = province
    if city:
        record["城市"] = city
    old_county = clean_text(record.get("区县"))
    invalid_county = (
        "税务" in old_county
        or "总局" in old_county
        or old_county.startswith(("家", "内到", "至", "向"))
        or old_county in {clean_text(record.get("省份")), clean_text(record.get("城市"))}
    )
    if county:
        record["区县"] = county
    elif invalid_county:
        record["区县"] = ""

    party = clean_text(record.get("当事人名称"))
    if not party.endswith(ENTITY_SUFFIXES):
        record["统一社会信用代码"] = ""
        record["法定代表人"] = ""
    legal_rep = clean_text(record.get("法定代表人"))
    if any(token in legal_rep for token in INVALID_LEGAL_REP_TOKENS):
        record["法定代表人"] = ""

    result = clean_text(record.get("处理或处罚结果"))
    if (
        clean_text(record.get("文书类型")) == "税务处理决定书"
        and "税务行政处罚决定书" in result
        and "罚款" in result
        and not any(token in result for token in ("追缴", "补缴"))
    ):
        record["处理或处罚结果"] = ""
        append_note(record, "处理文书中的处罚结果串值已根据官方附件清除。")

    # B 类记录只证明当事人、文书类型和文号。公告模板中的零散句子、税种和
    # 金额不能当作案情证据，统一留空，避免把半截文号或送达措辞展示成结果。
    if clean_text(record.get("公开完整度")) == "仅公告送达/仅文号线索":
        record["主要违法事实"] = ""
        record["涉及税种"] = ""
        record["处理或处罚结果"] = ""
        for field_name in MONEY_FIELDS:
            record[field_name] = None

    # 清理由换行/连续空白差异产生的历史伪冲突备注，不影响真实值冲突记录。
    note = clean_text(record.get("备注"))
    conflict_pattern = re.compile(
        r"字段“(?P<field>[^”]+)”存在冲突:保留原值“(?P<old>.*?)”[，,]"
        r"新值“(?P<new>.*?)”未静默覆盖。",
        re.S,
    )
    def keep_conflict(match: re.Match) -> str:
        field_name = match.group("field")
        old = match.group("old")
        new = match.group("new")
        if comparison_text(old) == comparison_text(new):
            return ""
        if field_name == "公开完整度" and old == "完整文书" and new == "仅公告送达/仅文号线索":
            return ""
        attachment_urls = {
            normalize_url(item)
            for item in clean_text(record.get("附件链接")).split(";")
            if clean_text(item)
        }
        if field_name == "官方原文链接" and normalize_url(new) in attachment_urls:
            return ""
        if field_name == "处理或处罚结果" and not clean_text(record.get("处理或处罚结果")):
            return ""
        if field_name in {"主要违法事实", "处理或处罚结果"} and f"字段“{field_name}”已根据官方附件重新解析。" in note:
            return ""
        if (
            field_name in {"主要违法事实", "处理或处罚结果"}
            and clean_text(record.get("来源级别")) == "税务机关官网"
            and clean_text(record.get("页面当前状态")) == "附件正常"
        ):
            return f"字段“{field_name}”历史冲突已由官方附件复核，保留当前值。"
        return match.group(0)

    note = conflict_pattern.sub(keep_conflict, note)
    record["备注"] = re.sub(r"\s{2,}", " ", note).strip()


def relink_case_groups(records: list[dict]) -> None:
    """统一同页同当事人的处理/处罚文书案件组，并重建双向关联文号。"""
    same_page: dict[tuple[str, str], list[dict]] = {}
    for record in records:
        record["案件组ID"] = clean_text(record.get("案件组ID")) or build_group_id(record)
        record["文书唯一ID"] = clean_text(record.get("文书唯一ID")) or build_unique_id(record)
        party = clean_text(record.get("当事人名称"))
        url = normalize_url(record.get("官方原文链接") or record.get("备用来源链接") or "")
        if party and url:
            same_page.setdefault((party, url), []).append(record)

    for page_records in same_page.values():
        if len({clean_text(item.get("文书类型")) for item in page_records}) < 2:
            continue
        canonical = min(
            page_records,
            key=lambda item: int(parse_money(item.get("序号")) or 10**9),
        )
        group_id = clean_text(canonical.get("案件组ID")) or build_group_id(canonical)
        for record in page_records:
            record["案件组ID"] = group_id

    groups: dict[str, list[dict]] = {}
    for record in records:
        doc_no = clean_text(record.get("决定书文号"))
        if record.get("文书类型") == "税务处理决定书":
            record["关联处理决定书文号"] = doc_no
            record["关联处罚决定书文号"] = ""
        else:
            record["关联处理决定书文号"] = ""
            record["关联处罚决定书文号"] = doc_no
        groups.setdefault(record["案件组ID"], []).append(record)

    for group_records in groups.values():
        processing = next((r.get("决定书文号", "") for r in group_records if r.get("文书类型") == "税务处理决定书"), "")
        penalty = next((r.get("决定书文号", "") for r in group_records if r.get("文书类型") == "税务行政处罚决定书"), "")
        for record in group_records:
            if processing:
                record["关联处理决定书文号"] = processing
            if penalty:
                record["关联处罚决定书文号"] = penalty


def merge_all(existing_xlsx: list[dict], existing_csv: list[dict], incoming: list[dict]) -> tuple[list[dict], dict]:
    records: list[dict] = []
    key_to_index: dict[tuple, int] = {}
    stats = {"new_full": 0, "new_clue": 0, "updated": 0, "duplicates": 0, "new_parties": []}

    def locate(record: dict) -> int | None:
        for key in canonical_keys(record):
            if key in key_to_index:
                return key_to_index[key]
        uid = clean_text(record.get("文书唯一ID"))
        if uid:
            for index, existing in enumerate(records):
                if clean_text(existing.get("文书唯一ID")) == uid:
                    return index
        return None

    def index_record(index: int, record: dict) -> None:
        for key in canonical_keys(record):
            key_to_index[key] = index

    # Excel 和 CSV 都必须读取；若两者不一致，取并集且保留冲突说明。
    for source_name, batch in (("Excel", existing_xlsx), ("CSV", existing_csv)):
        for record in batch:
            if provisional_out_of_scope(record):
                continue
            index = locate(record)
            if index is None:
                normalized = {field_name: record.get(field_name, "") for field_name in FIELDS}
                normalized["文书唯一ID"] = clean_text(normalized["文书唯一ID"]) or build_unique_id(normalized)
                normalized["案件组ID"] = clean_text(normalized["案件组ID"]) or build_group_id(normalized)
                records.append(normalized)
                index_record(len(records) - 1, normalized)
            else:
                merged, _ = merge_record(records[index], record)
                if source_name == "CSV":
                    append_note(merged, "Excel 与 CSV 历史数据已合并核对。")
                records[index] = merged
                index_record(index, merged)

    max_sequence = max(
        [int(parse_money(record.get("序号")) or 0) for record in records] or [0]
    )
    for record in incoming:
        if provisional_out_of_scope(record):
            continue
        index = locate(record)
        if index is None:
            max_sequence += 1
            record["序号"] = max_sequence
            record["文书唯一ID"] = record.get("文书唯一ID") or build_unique_id(record)
            record["案件组ID"] = record.get("案件组ID") or build_group_id(record)
            records.append(record)
            index_record(len(records) - 1, record)
            if record.get("公开完整度") == "完整文书":
                stats["new_full"] += 1
            else:
                stats["new_clue"] += 1
            party = clean_text(record.get("当事人名称"))
            if party and party not in stats["new_parties"]:
                stats["new_parties"].append(party)
        else:
            stats["duplicates"] += 1
            merged, changed = merge_record(records[index], record)
            records[index] = merged
            index_record(index, merged)
            if changed:
                stats["updated"] += 1

    for record in records:
        sanitize_record(record)
    # 两类文书仍保留独立记录；同一官方页面明确公开的成对文书统一案件组。
    relink_case_groups(records)

    def sort_key(record: dict) -> tuple:
        published = iso_date(record.get("官方发布日期")) or iso_date(record.get("第三方收录日期")) or "0000-00-00"
        inverse = tuple(-ord(char) for char in published)
        return (
            inverse,
            clean_text(record.get("省份")),
            clean_text(record.get("当事人名称")),
            clean_text(record.get("文书类型")),
        )

    records.sort(key=sort_key)
    return records, stats


def write_csv(path: Path, records: list[dict]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    # 临时文件仍保留 .xlsx 扩展名，确保 openpyxl 能重新打开执行替换前校验。
    temp = path.with_name(f"{path.stem}.tmp{path.suffix}")
    with temp.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=FIELDS, extrasaction="ignore", lineterminator="\n")
        writer.writeheader()
        for record in records:
            row = {}
            for field_name in FIELDS:
                value = record.get(field_name, "")
                if field_name in MONEY_FIELDS:
                    value = money_string(value)
                elif field_name in DATE_FIELDS:
                    value = iso_date(value)
                row[field_name] = value if value is not None else ""
            writer.writerow(row)
    os.replace(temp, path)


def style_sheet(sheet, headers: list[str], rows: list[dict], sheet_name: str) -> None:
    header_fill = PatternFill("solid", fgColor="17365D")
    header_font = Font(name="Microsoft YaHei", color="FFFFFF", bold=True, size=10)
    body_font = Font(name="Microsoft YaHei", size=9)
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    body_alignment = Alignment(vertical="top", wrap_text=True)

    sheet.sheet_view.showGridLines = False
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{max(1, len(rows) + 1)}"
    sheet.row_dimensions[1].height = 34
    for col_index, header in enumerate(headers, 1):
        cell = sheet.cell(1, col_index, header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = header_alignment

    hyperlink_labels = {"官方原文链接": "打开原文", "附件链接": "打开附件", "备用来源链接": "打开备用来源"}
    for row_index, record in enumerate(rows, 2):
        if sheet_name == "每日运行日志":
            sheet.row_dimensions[row_index].height = 54
        for col_index, header in enumerate(headers, 1):
            value = record.get(header, "")
            cell = sheet.cell(row_index, col_index)
            cell.font = body_font
            cell.alignment = body_alignment
            if header in MONEY_FIELDS:
                number = parse_money(value)
                cell.value = number
                cell.number_format = '#,##0.00'
                cell.alignment = Alignment(horizontal="right", vertical="top")
            elif header in DATE_FIELDS:
                normalized = iso_date(value)
                cell.value = date.fromisoformat(normalized) if normalized else None
                cell.number_format = "yyyy-mm-dd"
                cell.alignment = Alignment(horizontal="center", vertical="top")
            elif header in LINK_FIELDS:
                url = clean_text(value).split(";")[0]
                if url.startswith(("http://", "https://")):
                    cell.value = hyperlink_labels[header]
                    cell.hyperlink = url
                    cell.style = "Hyperlink"
                else:
                    cell.value = ""
            else:
                cell.value = excel_safe(value if value is not None else "")
            if header in LONG_TEXT_FIELDS:
                cell.alignment = Alignment(vertical="top", wrap_text=True)

    widths = {
        "序号": 8, "案件组ID": 22, "文书唯一ID": 22, "省份": 12, "城市": 13, "区县": 13,
        "发布机关": 34, "稽查机构": 30, "当事人名称": 30, "统一社会信用代码": 22,
        "法定代表人": 12, "文书类型": 23, "决定书文号": 28, "关联处理决定书文号": 28,
        "关联处罚决定书文号": 28, "主要违法事实": 58, "涉及税种": 28,
        "追缴税款金额": 16, "滞纳金金额": 16, "罚款金额": 16, "没收违法所得金额": 18,
        "处理或处罚结果": 55, "决定书作出日期": 15, "官方发布日期": 15,
        "第三方收录日期": 15, "首次发现日期": 15, "最后核验日期": 15,
        "公开完整度": 22, "来源级别": 17, "核验状态": 14, "官方原文链接": 14,
        "附件链接": 14, "备用来源链接": 16, "页面标题": 48, "页面当前状态": 16, "备注": 52,
    }
    for index, header in enumerate(headers, 1):
        sheet.column_dimensions[get_column_letter(index)].width = widths.get(header, 18)
    sheet.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{max(1, sheet.max_row)}"

    if headers == FIELDS and sheet.max_row >= 2:
        status_col = get_column_letter(headers.index("核验状态") + 1)
        page_col = get_column_letter(headers.index("页面当前状态") + 1)
        pending_fill = PatternFill("solid", fgColor="FFF2CC")
        bad_fill = PatternFill("solid", fgColor="F4CCCC")
        sheet.conditional_formatting.add(
            f"{status_col}2:{status_col}{sheet.max_row}",
            FormulaRule(formula=[f'${status_col}2="待核验"'], fill=pending_fill),
        )
        sheet.conditional_formatting.add(
            f"{page_col}2:{page_col}{sheet.max_row}",
            FormulaRule(formula=[f'AND(${page_col}2<>"正常",${page_col}2<>"附件正常")'], fill=bad_fill),
        )
        validation = DataValidation(
            type="list",
            formula1='"已核验,待核验,人工复核完成"',
            allow_blank=False,
        )
        sheet.add_data_validation(validation)
        validation.add(f"{status_col}2:{status_col}{max(2, sheet.max_row)}")

    sheet.sheet_properties.pageSetUpPr.fitToPage = True
    sheet.page_setup.fitToWidth = 1
    sheet.page_setup.fitToHeight = 0


def write_workbook(path: Path, records: list[dict], logs: list[dict]) -> None:
    workbook = Workbook()
    workbook.remove(workbook.active)
    subsets = {
        "文书汇总": records,
        "完整文书": [r for r in records if r.get("公开完整度") == "完整文书"],
        "公告送达及文号线索": [r for r in records if r.get("公开完整度") == "仅公告送达/仅文号线索"],
        "待核验记录": [r for r in records if r.get("核验状态") == "待核验"],
        "失效链接": [r for r in records if r.get("页面当前状态") not in NORMAL_STATES],
    }
    for sheet_name, rows in subsets.items():
        sheet = workbook.create_sheet(sheet_name)
        style_sheet(sheet, FIELDS, rows, sheet_name)

    compact_logs = []
    for log in logs:
        item = dict(log)
        error = clean_text(item.get("错误摘要"))
        if "baidu" in error.lower() and any(token in error.lower() for token in ("captcha", "timeout", "timed out")):
            item["错误摘要"] = "百度检索入口触发验证码或读取超时；其他检索入口继续执行，本次任务成功。"
        elif len(error) > 300:
            item["错误摘要"] = error[:297] + "..."
        parties = clean_text(item.get("本次新增当事人名单"))
        if len(parties) > 500:
            item["本次新增当事人名单"] = parties[:497] + "..."
        compact_logs.append(item)
    log_sheet = workbook.create_sheet("每日运行日志")
    style_sheet(log_sheet, LOG_FIELDS, compact_logs, "每日运行日志")
    log_widths = [22, 25, 14, 18, 18, 16, 16, 14, 16, 60, 14, 60, 60]
    for index, width in enumerate(log_widths, 1):
        log_sheet.column_dimensions[get_column_letter(index)].width = width

    workbook.calculation.fullCalcOnLoad = True
    workbook.calculation.forceFullCalc = True
    path.parent.mkdir(parents=True, exist_ok=True)
    temp = path.with_name(f"{path.stem}.tmp{path.suffix}")
    workbook.save(temp)
    workbook.close()
    # 在替换正式文件前重新打开临时文件，防止将损坏文件覆盖历史成品。
    check = load_workbook(temp, read_only=False, data_only=False)
    required = {"文书汇总", "完整文书", "公告送达及文号线索", "待核验记录", "失效链接", "每日运行日志"}
    if set(check.sheetnames) != required:
        check.close()
        raise RuntimeError(f"工作表校验失败：{check.sheetnames}")
    check.close()
    os.replace(temp, path)


def append_audit(entries: list[dict]) -> None:
    AUDIT_PATH.parent.mkdir(parents=True, exist_ok=True)
    with AUDIT_PATH.open("a", encoding="utf-8") as handle:
        for entry in entries:
            handle.write(json.dumps(entry, ensure_ascii=False, sort_keys=True) + "\n")


def read_json_object(path: Path, fallback: dict) -> dict:
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
        return value if isinstance(value, dict) else fallback
    except (OSError, ValueError, TypeError):
        return fallback


def atomic_write_json(path: Path, value: object) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temp = path.with_suffix(f".tmp{path.suffix}")
    temp.write_text(json.dumps(value, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    os.replace(temp, path)


def load_retry_queue() -> dict[str, dict]:
    payload = read_json_object(RETRY_QUEUE_PATH, {"items": []})
    items = payload.get("items", [])
    if not isinstance(items, list):
        return {}
    result: dict[str, dict] = {}
    for item in items:
        if not isinstance(item, dict):
            continue
        url = normalize_url(item.get("url", ""))
        if url and is_official_url(url):
            result[url] = {**item, "url": url}
    return result


def retry_queue_hits(queue: dict[str, dict]) -> list[SearchHit]:
    return [
        SearchHit(
            url=url,
            title=clean_text(item.get("title")),
            provider="持久化失败重试",
            query="retry-queue",
        )
        for url, item in sorted(queue.items())
    ]


def transient_access_failure(audit: dict) -> bool:
    status_code = audit.get("status_code")
    try:
        status_code = int(status_code) if status_code is not None else None
    except (TypeError, ValueError):
        status_code = None
    return bool(
        status_code in (403, 408, 412, 425, 429)
        or (status_code is not None and status_code >= 500)
        or (status_code is None and clean_text(audit.get("error")))
        or clean_text(audit.get("page_state")) == "暂时无法访问"
    )


def update_retry_queue(existing: dict[str, dict], audits: list[dict]) -> dict[str, dict]:
    queue = dict(existing)
    now = now_in_project_timezone()
    now_text = now.isoformat(timespec="seconds")
    for audit in audits:
        url = normalize_url(audit.get("url", ""))
        if not url or not is_official_url(url):
            continue
        # 已打开但解析失败的目标文书也要保留，规则修复后才能再次处理。
        parse_failure = audit.get("skip_reason") == "无法确认当事人"
        if transient_access_failure(audit) or parse_failure:
            previous = queue.get(url, {})
            failures = int(previous.get("failureCount", 0) or 0) + 1
            delay_hours = 1 if failures == 1 else 6 if failures == 2 else 24
            queue[url] = {
                "url": url,
                "title": clean_text(audit.get("title")) or clean_text(previous.get("title")),
                "provider": clean_text(audit.get("provider")) or clean_text(previous.get("provider")),
                "query": clean_text(audit.get("query")) or clean_text(previous.get("query")),
                "firstFailedAt": previous.get("firstFailedAt") or now_text,
                "lastFailedAt": now_text,
                "failureCount": failures,
                "nextRetryAt": (now + timedelta(hours=delay_hours)).isoformat(timespec="seconds"),
                "lastError": clean_text(audit.get("error")) or clean_text(audit.get("skip_reason")) or clean_text(audit.get("page_state")),
                "lastStatus": audit.get("status_code"),
            }
        else:
            # 已恢复、明确 404/410 或内容可访问时移出临时失败队列；历史记录仍会按状态每日复核。
            queue.pop(url, None)
    atomic_write_json(RETRY_QUEUE_PATH, {
        "updatedAt": now_text,
        "items": sorted(queue.values(), key=lambda item: item["url"]),
    })
    return queue


def province_for_audit(audit: dict) -> str:
    url = normalize_url(audit.get("final_url") or audit.get("url") or "")
    host_prefix = urlparse(url).netloc.lower().split(".")[0]
    if host_prefix in DOMAIN_PROVINCES:
        return DOMAIN_PROVINCES[host_prefix]
    province, _, _ = locate_region(clean_text(audit.get("title")), "", url)
    return province


def write_source_health(audits: list[dict]) -> tuple[dict, list[str]]:
    previous_payload = read_json_object(PUBLIC_SOURCE_HEALTH_PATH, {"sources": []})
    previous_sources = {
        item.get("source"): item
        for item in previous_payload.get("sources", [])
        if isinstance(item, dict) and item.get("source")
    }
    now_text = now_in_project_timezone().isoformat(timespec="seconds")
    sources = []
    failed_sources = []
    retry_queue = load_retry_queue()
    for province in CORE_SOURCE_PROVINCES:
        matching = [audit for audit in audits if province_for_audit(audit) == province]
        successful = sum(
            1 for audit in matching
            if not transient_access_failure(audit)
            and audit.get("status_code") is not None
            and 200 <= int(audit["status_code"]) < 500
        )
        failed = sum(1 for audit in matching if transient_access_failure(audit))
        if matching:
            health_status = "healthy" if failed == 0 else "degraded" if successful else "unreachable"
            item = {
                "source": province,
                "checked": len(matching),
                "checkedThisRun": len(matching),
                "succeeded": successful,
                "failed": failed,
                "status": health_status,
                "lastCheckedAt": now_text,
                "lastError": next(
                    (clean_text(audit.get("error")) or clean_text(audit.get("page_state")) for audit in matching if transient_access_failure(audit)),
                    "",
                ),
            }
        else:
            old = previous_sources.get(province, {})
            item = {
                "source": province,
                "checked": old.get("checked", 0),
                "checkedThisRun": 0,
                "succeeded": old.get("succeeded", 0),
                "failed": old.get("failed", 0),
                "status": old.get("status", "not_checked"),
                "lastCheckedAt": old.get("lastCheckedAt"),
                "lastError": old.get("lastError", ""),
            }
        item["pendingRetries"] = sum(1 for url in retry_queue if province_for_audit({"url": url}) == province)
        if item["pendingRetries"] and item["status"] == "healthy":
            item["status"] = "degraded"
        if item["status"] in {"degraded", "unreachable"}:
            failed_sources.append(province)
        sources.append(item)
    for audit in audits:
        if not transient_access_failure(audit):
            continue
        url = normalize_url(audit.get("final_url") or audit.get("url") or "")
        source_name = province_for_audit(audit) or urlparse(url).netloc.lower()
        if source_name and source_name not in failed_sources:
            failed_sources.append(source_name)
    payload = {"generatedAt": now_text, "sources": sources}
    atomic_write_json(PUBLIC_SOURCE_HEALTH_PATH, payload)
    return payload, failed_sources


def current_source_commit() -> str:
    configured = clean_text(os.getenv("SOURCE_COMMIT") or os.getenv("GITHUB_SHA"))
    if configured:
        return configured
    try:
        result = subprocess.run(
            ["git", "rev-parse", "HEAD"],
            cwd=ROOT,
            check=True,
            capture_output=True,
            text=True,
            timeout=5,
        )
        return clean_text(result.stdout)
    except (OSError, subprocess.SubprocessError):
        return "local"


def public_record(record: dict) -> dict:
    item: dict[str, object] = {}
    for json_name, source_name in JSON_FIELD_MAP.items():
        value = record.get(source_name)
        if source_name in MONEY_FIELDS:
            item[json_name] = parse_money(value)
        elif source_name in DATE_FIELDS:
            item[json_name] = iso_date(value) or None
        else:
            cleaned = clean_text(value)
            item[json_name] = cleaned or None
    return item


def build_search_coverage(audits: list[dict], queue: dict[str, dict], run_mode: str, records: list[dict] | None = None) -> dict:
    """任务写入成功不等于官网覆盖完整；失败候选与正式文书分开公布。"""
    accessible = sum(
        1 for audit in audits
        if isinstance(audit.get("status_code"), int) and 200 <= audit["status_code"] < 300
        and audit.get("page_state") in NORMAL_STATES
    )
    failed = sum(1 for audit in audits if transient_access_failure(audit))
    unresolved = sum(1 for audit in audits if audit.get("skip_reason") == "无法确认当事人")
    pending = []
    collected_urls = {
        normalize_url(url)
        for record in records or []
        for field in ("官方原文链接", "附件链接", "备用来源链接")
        for url in clean_text(record.get(field)).split(";") if url
    }
    cutoff = TODAY - timedelta(days=OFFICIAL_INDEX_LOOKBACK_DAYS)
    for url, item in queue.items():
        hint = date_from_listing_url(url)
        if url in collected_urls or not hint or not cutoff <= hint <= TODAY:
            continue
        pending.append({
            "url": url,
            "title": clean_text(item.get("title")) or "官方候选页面（标题待核验）",
            "dateHint": hint.isoformat(),
            "lastCheckedAt": item.get("lastFailedAt"),
            "reason": clean_text(item.get("lastError")),
        })
    pending.sort(key=lambda item: (item["dateHint"], item["url"]), reverse=True)
    return {
        "runMode": run_mode,
        "coverageStatus": "unavailable" if not accessible else (
            "partial" if failed or unresolved or queue or run_mode != "scheduled" else "complete"
        ),
        "auditedPages": len(audits),
        "accessiblePages": accessible,
        "failedPages": failed,
        "unresolvedPages": unresolved,
        "retryQueueSize": len(queue),
        "pendingCandidates": pending,
    }


def write_public_artifacts(records: list[dict], log_entry: dict, failed_sources: list[str], coverage: dict | None = None) -> None:
    PUBLIC_DATA_DIR.mkdir(parents=True, exist_ok=True)
    PUBLIC_DOWNLOAD_DIR.mkdir(parents=True, exist_ok=True)
    public_records = [public_record(record) for record in records]
    temp_json = PUBLIC_JSON_PATH.with_suffix(".tmp.json")
    temp_json.write_text(
        json.dumps(public_records, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    today_new = sum(
        1 for record in public_records
        if record.get("firstDiscoveredDate") == TODAY.isoformat()
    )
    previous_status = read_json_object(PUBLIC_STATUS_PATH, {})
    completed_at = now_in_project_timezone().isoformat(timespec="seconds")
    run_success = log_entry["运行是否成功"] == "成功"
    data_changed = (
        int(log_entry["新增完整文书数量"])
        + int(log_entry["新增文号线索数量"])
        + int(log_entry["更新旧记录数量"])
    ) > 0
    status = {
        "status": "normal" if run_success and not failed_sources and (coverage or {}).get("coverageStatus", "complete") == "complete" else "degraded",
        "lastUpdated": completed_at,
        "lastRunStartedAt": log_entry["运行日期和时间"],
        "lastRunCompletedAt": completed_at,
        "lastSuccessfulRunAt": completed_at if run_success else previous_status.get("lastSuccessfulRunAt"),
        "lastDataChangeAt": completed_at if data_changed else previous_status.get("lastDataChangeAt", completed_at),
        "lastProductionDeploymentAt": previous_status.get("lastProductionDeploymentAt"),
        "sourceCommit": current_source_commit(),
        "workflowRunId": os.getenv("GITHUB_RUN_ID") or f"local-{RUN_NOW.strftime('%Y%m%d%H%M%S')}",
        "deploymentStatus": "pending" if os.getenv("GITHUB_ACTIONS") == "true" else "local",
        "timezone": TZ_NAME,
        "nextScheduledUpdate": PUBLIC_SCHEDULE,
        "searchRange": log_entry["检索时间范围"],
        "searchedPages": log_entry["检索页面数量"],
        "total": len(public_records),
        "todayNew": today_new,
        "completeDocuments": sum(1 for record in public_records if record.get("completeness") == "完整文书"),
        "clues": sum(1 for record in public_records if record.get("completeness") != "完整文书"),
        "pending": sum(1 for record in public_records if record.get("verificationStatus") == "待核验"),
        "invalidLinks": sum(1 for record in public_records if record.get("pageStatus") not in NORMAL_STATES),
        "newCompleteDocuments": log_entry["新增完整文书数量"],
        "newClues": log_entry["新增文号线索数量"],
        "newRecords": int(log_entry["新增完整文书数量"]) + int(log_entry["新增文号线索数量"]),
        "updatedRecords": log_entry["更新旧记录数量"],
        "duplicateRecords": log_entry["重复记录数量"],
        "failedSources": failed_sources,
        "runSuccess": run_success,
        "message": log_entry["运行说明"],
        **(coverage or {}),
    }
    os.replace(temp_json, PUBLIC_JSON_PATH)
    atomic_write_json(PUBLIC_STATUS_PATH, status)
    shutil.copy2(XLSX_PATH, PUBLIC_XLSX_PATH)
    shutil.copy2(CSV_PATH, PUBLIC_CSV_PATH)


def process_hits(hits: list[SearchHit], start_date: date, full_run: bool) -> tuple[list[dict], list[dict]]:
    records: list[dict] = []
    audits: list[dict] = []
    console(f"核验候选页面：{len(hits)} 个")
    with ThreadPoolExecutor(max_workers=8) as executor:
        future_map = {executor.submit(fetch_candidate, hit): hit for hit in hits}
        for index, future in enumerate(as_completed(future_map), 1):
            hit = future_map[future]
            try:
                fetched = future.result()
                parsed, audit = parse_fetch(hit, fetched, start_date, full_run)
                records.extend(parsed)
                audits.append(audit)
            except Exception as exc:
                audits.append({
                    "checked_at": RUN_NOW.isoformat(timespec="seconds"),
                    "url": hit.url,
                    "provider": hit.provider,
                    "query": hit.query,
                    "page_state": "需要人工核验",
                    "error": f"{type(exc).__name__}: {exc}",
                    "records": 0,
                })
            if index % 25 == 0:
                console(f"已核验 {index}/{len(hits)} 个候选页面")
    return records, audits


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--force-full", action="store_true", help="强制从 2026-01-01 全量回溯，用于审计/幂等测试")
    parser.add_argument("--start-date", help="覆盖检索起始日期，格式 YYYY-MM-DD")
    parser.add_argument("--max-pages", type=int, default=0, help="调试用：限制核验页面数；0 表示不限制")
    parser.add_argument("--urls-file", type=Path, help="定向补录：逐条联网核验文件内的官方链接，保留历史数据与未处理重试队列")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    existing_xlsx, logs = read_xlsx_records(XLSX_PATH)
    existing_csv = read_csv_records(CSV_PATH)
    first_run = not existing_xlsx and not existing_csv
    full_run = first_run or args.force_full
    if args.start_date:
        start_date = date.fromisoformat(args.start_date)
    elif full_run:
        start_date = date(2026, 1, 1)
    else:
        start_date = TODAY - timedelta(days=RECENT_LOOKBACK_DAYS)
    if start_date < date(2026, 1, 1):
        start_date = date(2026, 1, 1)

    errors: list[str] = []
    success = True
    stats = {"new_full": 0, "new_clue": 0, "updated": 0, "duplicates": 0, "new_parties": []}
    audits: list[dict] = []
    discovery_audits: list[dict] = []
    records: list[dict] = []
    hits: list[SearchHit] = []
    existing_retry_queue = load_retry_queue()
    console(f"检索时间范围：{start_date.isoformat()} 至 {TODAY.isoformat()}（{TZ_NAME}）")
    try:
        if args.urls_file:
            urls = [normalize_url(line.strip()) for line in args.urls_file.read_text(encoding="utf-8-sig").splitlines()
                    if line.strip() and not line.lstrip().startswith("#")]
            if not urls or any(not is_official_url(url) for url in urls):
                raise ValueError("定向补录文件必须包含有效的官方链接")
            hits = [SearchHit(url=url, provider="定向官方补录", query="retry-queue") for url in dict.fromkeys(urls)]
        else:
            hits = discover_candidates(full_run, errors, existing_xlsx or existing_csv, discovery_audits)
            historical_hits = historical_revalidation_hits(existing_xlsx or existing_csv)
            queued_hits = retry_queue_hits(existing_retry_queue)
            # 固定优先级：持久化失败/历史复核 > 官方栏目 > 搜索引擎。
            hits = merge_search_hits(hits, historical_hits, queued_hits)
        if args.max_pages:
            hits = hits[:args.max_pages]
        console(f"发现并去重候选链接：{len(hits)} 个")
        incoming, page_audits = process_hits(hits, start_date, full_run)
        audits = discovery_audits + page_audits
        records, stats = merge_all(existing_xlsx, existing_csv, incoming)
    except Exception as exc:
        success = False
        errors.append(f"主流程：{type(exc).__name__}: {exc}")
        records, _ = merge_all(existing_xlsx, existing_csv, [])

    if not audits:
        audits = discovery_audits
    append_audit(audits)
    try:
        retry_queue = update_retry_queue(existing_retry_queue, audits)
        _, failed_sources = write_source_health(audits)
    except Exception as exc:
        console(f"重试队列/来源健康状态写入失败：{type(exc).__name__}: {exc}")
        return 1
    failed_count = sum(1 for record in records if record.get("页面当前状态") not in NORMAL_STATES)
    pending_count = sum(1 for record in records if record.get("核验状态") == "待核验")
    coverage = build_search_coverage(audits, retry_queue, "targeted" if args.urls_file or args.max_pages else "scheduled", records)
    if stats["new_full"] + stats["new_clue"] == 0:
        run_note = "本次检索未发现符合收录标准的新文书。"
    else:
        run_note = "已完成候选发现、逐条访问核验、去重、关联和历史数据增量写入。"
    if coverage["coverageStatus"] != "complete":
        run_note = (
            f"本次已核验来源新增{stats['new_full'] + stats['new_clue']}份文书；检索覆盖不完整，不能据此认定其他日期没有新文书。"
            f"本轮可访问{coverage['accessiblePages']}页，访问失败{coverage['failedPages']}页，持久化待复核{coverage['retryQueueSize']}个链接。"
            + ("本次为定向补录，未执行全国全量检索。" if args.urls_file else "")
        )
    log_entry = {
        "运行日期和时间": RUN_NOW.isoformat(timespec="seconds"),
        "检索时间范围": f"{start_date.isoformat()} 至 {TODAY.isoformat()}",
        "检索页面数量": len(hits),
        "新增完整文书数量": stats["new_full"],
        "新增文号线索数量": stats["new_clue"],
        "更新旧记录数量": stats["updated"],
        "重复记录数量": stats["duplicates"],
        "失效链接数量": failed_count,
        "待人工核验数量": pending_count,
        "本次新增当事人名单": "、".join(stats["new_parties"]),
        "运行是否成功": "成功" if success else "失败",
        "错误摘要": " | ".join(errors[:20]),
        "运行说明": run_note,
    }
    logs.append(log_entry)

    try:
        write_csv(CSV_PATH, records)
        write_workbook(XLSX_PATH, records, logs)
        write_public_artifacts(records, log_entry, failed_sources, coverage)
    except Exception as exc:
        console(f"输出失败：{type(exc).__name__}: {exc}")
        return 1

    console(json.dumps({
        "success": success,
        "range": log_entry["检索时间范围"],
        "pages": len(hits),
        "total_records": len(records),
        "new_full": stats["new_full"],
        "new_clue": stats["new_clue"],
        "updated": stats["updated"],
        "duplicates": stats["duplicates"],
        "pending": pending_count,
        "invalid_links": failed_count,
        "xlsx": str(XLSX_PATH),
        "csv": str(CSV_PATH),
    }, ensure_ascii=False, indent=2))
    return 0 if success else 2


if __name__ == "__main__":
    sys.exit(main())
